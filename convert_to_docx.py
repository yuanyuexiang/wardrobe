#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将用户操作手册从Markdown格式转换为DOCX格式
特别处理emoji符号和中文内容
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_styled_document():
    """创建带有样式的文档"""
    doc = Document()
    
    # 设置文档的默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'  # 支持中文和emoji的字体
    font.size = Pt(11)
    
    # 创建标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Microsoft YaHei'
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(18)
        elif i == 2:
            heading_font.size = Pt(16)
        else:
            heading_font.size = Pt(14)
    
    return doc

def parse_markdown_line(line):
    """解析markdown行，返回样式和内容"""
    line = line.strip()
    
    if not line:
        return 'normal', ''
    
    # 标题
    if line.startswith('# '):
        return 'heading1', line[2:]
    elif line.startswith('## '):
        return 'heading2', line[3:]
    elif line.startswith('### '):
        return 'heading3', line[4:]
    elif line.startswith('#### '):
        return 'heading4', line[5:]
    
    # 列表项
    elif line.startswith('- '):
        return 'bullet', line[2:]
    elif re.match(r'^\d+\. ', line):
        return 'numbered', re.sub(r'^\d+\. ', '', line)
    
    # 粗体文本
    elif line.startswith('**') and line.endswith('**'):
        return 'bold', line[2:-2]
    
    # 普通文本
    else:
        return 'normal', line

def add_paragraph_with_style(doc, style, text):
    """添加带样式的段落"""
    if style == 'heading1':
        p = doc.add_heading(text, level=1)
    elif style == 'heading2':
        p = doc.add_heading(text, level=2)
    elif style == 'heading3':
        p = doc.add_heading(text, level=3)
    elif style == 'heading4':
        p = doc.add_heading(text, level=4)
    elif style == 'bullet':
        p = doc.add_paragraph(text, style='List Bullet')
    elif style == 'numbered':
        p = doc.add_paragraph(text, style='List Number')
    elif style == 'bold':
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
    else:
        p = doc.add_paragraph(text)
    
    # 确保字体支持emoji
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """将Markdown文件转换为DOCX"""
    doc = create_styled_document()
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_list_level = 0
    
    for line in lines:
        style, content = parse_markdown_line(line)
        
        if content:  # 只处理非空内容
            add_paragraph_with_style(doc, style, content)
        elif style == 'normal' and not content:
            # 添加空行
            doc.add_paragraph()
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"✅ 成功转换文档: {docx_file_path}")

def main():
    """主函数"""
    md_file = "用户操作手册.md"
    docx_file = "衣橱APP用户操作手册.docx"
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print(f"📄 文档已生成: {docx_file}")
        print("🎉 转换完成！emoji符号已正确处理。")
    except Exception as e:
        print(f"❌ 转换失败: {str(e)}")

if __name__ == "__main__":
    main()