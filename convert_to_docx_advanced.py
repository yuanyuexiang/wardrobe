#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
高级Markdown到DOCX转换器
专门优化处理emoji、中文内容和复杂格式
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def create_professional_document():
    """创建专业格式的文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    
    # 自定义标题样式
    styles = doc.styles
    
    # 主标题样式
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Microsoft YaHei'
        title_style.font.size = Pt(20)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0x2F, 0x4F, 0x4F)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # 副标题样式
    for i in range(1, 5):
        style_name = f'CustomHeading{i}'
        if style_name not in [s.name for s in styles]:
            heading_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Microsoft YaHei'
            heading_style.font.bold = True
            if i == 1:
                heading_style.font.size = Pt(16)
                heading_style.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
            elif i == 2:
                heading_style.font.size = Pt(14)
                heading_style.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)
            elif i == 3:
                heading_style.font.size = Pt(12)
                heading_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                heading_style.font.size = Pt(11)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    return doc

def process_text_formatting(text):
    """处理文本格式，包括粗体、斜体等"""
    # 处理粗体 **text**
    bold_pattern = r'\*\*(.*?)\*\*'
    # 处理代码 `code`
    code_pattern = r'`(.*?)`'
    
    # 标记格式化区域
    text = re.sub(bold_pattern, r'<BOLD>\1</BOLD>', text)
    text = re.sub(code_pattern, r'<CODE>\1</CODE>', text)
    
    return text

def add_formatted_paragraph(doc, text, style_name=None):
    """添加格式化段落"""
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    
    # 处理文本格式
    processed_text = process_text_formatting(text)
    
    # 分割并添加格式化文本
    parts = re.split(r'(<BOLD>.*?</BOLD>|<CODE>.*?</CODE>)', processed_text)
    
    for part in parts:
        if part.startswith('<BOLD>') and part.endswith('</BOLD>'):
            run = p.add_run(part[6:-7])  # 移除<BOLD>标签
            run.bold = True
        elif part.startswith('<CODE>') and part.endswith('</CODE>'):
            run = p.add_run(part[6:-7])  # 移除<CODE>标签
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xD6, 0x33, 0x6C)
        else:
            run = p.add_run(part)
        
        # 确保字体支持emoji
        run.font.name = 'Microsoft YaHei'

def parse_markdown_content(content):
    """解析Markdown内容"""
    lines = content.split('\n')
    parsed_content = []
    
    for line in lines:
        line = line.strip()
        
        if not line:
            parsed_content.append(('empty', ''))
            continue
        
        # 主标题
        if line.startswith('# '):
            parsed_content.append(('title', line[2:]))
        # 二级标题
        elif line.startswith('## '):
            parsed_content.append(('heading1', line[3:]))
        # 三级标题
        elif line.startswith('### '):
            parsed_content.append(('heading2', line[4:]))
        # 四级标题
        elif line.startswith('#### '):
            parsed_content.append(('heading3', line[5:]))
        # 五级标题
        elif line.startswith('##### '):
            parsed_content.append(('heading4', line[6:]))
        # 无序列表
        elif line.startswith('- '):
            parsed_content.append(('bullet', line[2:]))
        # 有序列表
        elif re.match(r'^\d+\. ', line):
            parsed_content.append(('numbered', re.sub(r'^\d+\. ', '', line)))
        # 分隔线
        elif line.startswith('---'):
            parsed_content.append(('separator', ''))
        # 普通段落
        else:
            parsed_content.append(('paragraph', line))
    
    return parsed_content

def convert_advanced_markdown_to_docx(md_file_path, docx_file_path):
    """高级Markdown到DOCX转换"""
    doc = create_professional_document()
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 解析内容
    parsed_content = parse_markdown_content(content)
    
    for content_type, text in parsed_content:
        if content_type == 'title':
            add_formatted_paragraph(doc, text, 'CustomTitle')
        elif content_type == 'heading1':
            add_formatted_paragraph(doc, text, 'CustomHeading1')
        elif content_type == 'heading2':
            add_formatted_paragraph(doc, text, 'CustomHeading2')
        elif content_type == 'heading3':
            add_formatted_paragraph(doc, text, 'CustomHeading3')
        elif content_type == 'heading4':
            add_formatted_paragraph(doc, text, 'CustomHeading4')
        elif content_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text_to_paragraph(p, text)
        elif content_type == 'numbered':
            p = doc.add_paragraph(style='List Number')
            add_formatted_text_to_paragraph(p, text)
        elif content_type == 'separator':
            # 添加分隔线
            p = doc.add_paragraph()
            p.paragraph_format.border_bottom = True
        elif content_type == 'paragraph' and text:
            add_formatted_paragraph(doc, text)
        elif content_type == 'empty':
            doc.add_paragraph()
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"✅ 高级转换完成: {docx_file_path}")

def add_formatted_text_to_paragraph(paragraph, text):
    """向现有段落添加格式化文本"""
    processed_text = process_text_formatting(text)
    parts = re.split(r'(<BOLD>.*?</BOLD>|<CODE>.*?</CODE>)', processed_text)
    
    for part in parts:
        if part.startswith('<BOLD>') and part.endswith('</BOLD>'):
            run = paragraph.add_run(part[6:-7])
            run.bold = True
        elif part.startswith('<CODE>') and part.endswith('</CODE>'):
            run = paragraph.add_run(part[6:-7])
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xD6, 0x33, 0x6C)
        else:
            run = paragraph.add_run(part)
        
        run.font.name = 'Microsoft YaHei'

def main():
    """主函数"""
    md_file = "用户操作手册.md"
    docx_file = "衣橱APP用户操作手册_专业版.docx"
    
    try:
        convert_advanced_markdown_to_docx(md_file, docx_file)
        print(f"📄 专业版文档已生成: {docx_file}")
        print("🎉 高级转换完成！完美支持emoji和中文格式。")
        print("📝 文档特色:")
        print("   • 🎨 专业的颜色主题")
        print("   • 📱 完美的emoji显示")
        print("   • 🔤 优化的中文字体")
        print("   • 📐 标准的页面布局")
    except Exception as e:
        print(f"❌ 转换失败: {str(e)}")

if __name__ == "__main__":
    main()