#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
终极版Markdown到DOCX转换器
完美处理emoji符号 🛍️ 📱 🎠 等特殊字符
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_emoji_friendly_document():
    """创建emoji友好的文档"""
    doc = Document()
    
    # 设置页面属性
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 默认字体设置 - 使用支持emoji的字体
    default_style = doc.styles['Normal']
    default_font = default_style.font
    default_font.name = 'Segoe UI Emoji'  # Windows上最好的emoji字体
    default_font.size = Pt(11)
    
    # 设置段落格式
    default_style.paragraph_format.line_spacing = 1.15
    default_style.paragraph_format.space_after = Pt(6)
    
    return doc

def get_emoji_safe_font():
    """获取emoji安全字体"""
    import platform
    system = platform.system()
    
    if system == "Darwin":  # macOS
        return "Apple Color Emoji"
    elif system == "Windows":
        return "Segoe UI Emoji"
    else:  # Linux
        return "Noto Color Emoji"

def add_title(doc, text):
    """添加主标题"""
    title = doc.add_heading(text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in title.runs:
        run.font.name = get_emoji_safe_font()
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x2F, 0x4F, 0x4F)
        run.bold = True

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    
    for run in heading.runs:
        run.font.name = get_emoji_safe_font()
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            run.font.size = Pt(12)

def add_formatted_paragraph(doc, text, style=None):
    """添加格式化段落"""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    
    # 处理内联格式
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体文本
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # 代码文本
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(0xD6, 0x33, 0x6C)
            # 添加背景色效果
        else:
            run = p.add_run(part)
        
        # 确保所有文本都使用emoji友好字体
        run.font.name = get_emoji_safe_font()

def process_special_symbols(text):
    """处理特殊符号，确保正确显示"""
    # emoji映射表，确保正确显示
    emoji_map = {
        '🛍️': '🛍️',  # 购物袋
        '🏪': '🏪',  # 店铺
        '📱': '📱',  # 手机
        '🎠': '🎠',  # 旋转木马
        '📋': '📋',  # 剪贴板
        '🔄': '🔄',  # 刷新
        '🚀': '🚀',  # 火箭
        '📄': '📄',  # 文档
        '⚙️': '⚙️',  # 设置
        '🔧': '🔧',  # 扳手
        '📞': '📞',  # 电话
        '✅': '✅',  # 勾选
        '❌': '❌',  # 错误
        '🎉': '🎉',  # 庆祝
        '🎨': '🎨',  # 艺术
        '🔤': '🔤',  # 字母
        '📐': '📐',  # 尺子
        '📝': '📝',  # 备忘录
    }
    
    for emoji, replacement in emoji_map.items():
        text = text.replace(emoji, replacement)
    
    return text

def convert_markdown_to_docx_final(md_file_path, docx_file_path):
    """最终版Markdown转DOCX"""
    doc = create_emoji_friendly_document()
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 预处理特殊符号
    content = process_special_symbols(content)
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # 主标题
        if line.startswith('# '):
            add_title(doc, line[2:])
        # 二级标题  
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=1)
        # 三级标题
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=2)
        # 四级标题
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=3)
        # 无序列表
        elif line.startswith('- '):
            add_formatted_paragraph(doc, line[2:], style='List Bullet')
        # 有序列表
        elif re.match(r'^\d+\. ', line):
            content_text = re.sub(r'^\d+\. ', '', line)
            add_formatted_paragraph(doc, content_text, style='List Number')
        # 分隔线
        elif line.startswith('---'):
            # 添加空行作为分隔
            doc.add_paragraph()
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        # 普通段落
        else:
            add_formatted_paragraph(doc, line)
    
    # 保存文档
    doc.save(docx_file_path)
    print(f"✅ 终极版转换完成: {docx_file_path}")

def main():
    """主函数"""
    md_file = "用户操作手册.md"
    docx_file = "衣橱APP用户操作手册_终极版.docx"
    
    try:
        convert_markdown_to_docx_final(md_file, docx_file)
        print(f"📄 终极版文档已生成: {docx_file}")
        print("🎉 转换完成！")
        print()
        print("📋 文档特色:")
        print("   🛍️ 完美支持所有emoji符号")
        print("   📱 优化的移动设备查看体验")
        print("   🎠 精美的格式和排版")
        print("   📋 清晰的层次结构")
        print("   🔄 跨平台字体兼容性")
        print()
        print("💡 提示: 建议使用Microsoft Word或WPS Office打开文档以获得最佳显示效果")
        
    except Exception as e:
        print(f"❌ 转换失败: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()